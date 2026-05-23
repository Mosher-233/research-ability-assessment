#!/usr/bin/env python3
"""Generate BUPT thesis .docx with full formatting compliance.
Reads real project code for appendix, applies all school formatting rules."""

import os, re, io, copy
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DRAFT = os.path.join(BASE, "papers", "draft")
FIG   = os.path.join(BASE, "papers", "figures")
OUT   = os.path.join(BASE, "output", "doc")
os.makedirs(OUT, exist_ok=True)

# ── Read project source code for appendix (trimmed to core methods) ──
def read_src(rel):
    p = os.path.join(BASE, rel)
    if os.path.exists(p):
        with open(p, 'r', encoding='utf-8') as f:
            return f.read()
    return f"// file not found: {rel}"

def trim_go(src, keep_funcs):
    """Extract only specified function/method blocks from Go source."""
    lines = src.split('\n')
    result = []
    in_keep = False
    brace_depth = 0
    for line in lines:
        stripped = line.strip()
        # Check if line starts a function we want to keep
        for kf in keep_funcs:
            if kf in stripped and ('func ' in stripped or stripped.startswith('type ')):
                in_keep = True
                brace_depth = 0
                break
        if in_keep:
            result.append(line)
            brace_depth += stripped.count('{') - stripped.count('}')
            if brace_depth <= 0 and '{' in stripped:
                in_keep = False
                # Add a blank separator
                if result and result[-1] != '':
                    result.append('')
    if not result:
        return src  # fallback: return full file
    return '\n'.join(result)

# Load full source then trim to core methods
src_control_unit_full = read_src("internal/agent/control_unit.go")
src_control_unit = trim_go(src_control_unit_full,
    ['func (c *ControlUnit) ExecuteEvaluation', 'func (c *ControlUnit) NewControlUnit'])

src_evidence_full = read_src("internal/agent/evidence_agent.go")
src_evidence = trim_go(src_evidence_full,
    ['func (a *EvidenceAgent) classifyWithLLM', 'func (a *EvidenceAgent) ClassifyEvidence',
     'func (a *EvidenceAgent) extractKBMWithLLM', 'func (a *EvidenceAgent) ExtractKBMInfo'])

src_inference_full = read_src("internal/agent/inference_agent.go")
src_inference = trim_go(src_inference_full,
    ['func (a *InferenceAgent) InferAbility', 'func (a *InferenceAgent) inferWithLLM',
     'func (a *InferenceAgent) buildInferenceSystemPrompt', 'func (a *InferenceAgent) parseLLMInferenceResponse',
     'func (a *InferenceAgent) inferWithRules'])

src_feedback_full = read_src("internal/agent/feedback_agent.go")
src_feedback = trim_go(src_feedback_full,
    ['func (a *FeedbackAgent) generateLLMFeedback', 'func (a *FeedbackAgent) GenerateFeedback',
     'func (a *FeedbackAgent) generateRuleFeedback'])

src_logic        = read_src("internal/agent/logic_unit.go")
src_extractor    = read_src("pkg/extractor/extractor.go")
src_citation     = read_src("internal/models/citation.go")

src_auth_full    = read_src("internal/service/auth_service.go")
src_auth = trim_go(src_auth_full,
    ['func (s *AuthService) Login', 'func (s *AuthService) Register'])

src_llm_client   = read_src("internal/llm/client.go")

doc = Document()

# ═══════════════════ GLOBAL STYLE SETUP ═══════════════════
style_n = doc.styles['Normal']
style_n.font.name = '宋体'
style_n.font.size = Pt(12)
style_n.font.color.rgb = RGBColor(0, 0, 0)
style_n.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style_n.paragraph_format
pf.line_spacing = 1.5
pf.space_before = Pt(0)
pf.space_after = Pt(0)

def _set_hs(name, ea, ascii_f, sz, bld, al, ls, sb, sa, fi_chars=0):
    try: hs = doc.styles[name]
    except: hs = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    hs.font.name = ascii_f
    hs.font.size = Pt(sz)
    hs.font.bold = bld
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), ea)
    hpf = hs.paragraph_format
    hpf.alignment = al
    hpf.line_spacing = ls
    hpf.space_before = Pt(sb)
    hpf.space_after = Pt(sa)
    if fi_chars: hpf.first_line_indent = Pt(fi_chars * 12)
    return hs

# H1: 黑体三号(16pt) bold center, 段前0段后2行
_set_hs('Heading 1', '黑体', 'Times New Roman', 16, True, WD_ALIGN_PARAGRAPH.CENTER, 1.5, 0, 24)
# H2: 黑体四号(14pt) bold left, 段前段后0.5行  
_set_hs('Heading 2', '黑体', 'Times New Roman', 14, True, WD_ALIGN_PARAGRAPH.LEFT, 1.5, 6, 6)
# H3: 黑体小四(12pt) bold left, 首行缩进2字符, 段前段后0.5行
_set_hs('Heading 3', '黑体', 'Times New Roman', 12, True, WD_ALIGN_PARAGRAPH.LEFT, 1.5, 6, 6, fi_chars=2)

# ── Helper: run with fonts ──
def mk_run(p, text, ea='宋体', ascii_f='Times New Roman', sz=12, bold=False, italic=False, underline=False, color=None, superscript=False):
    r = p.add_run(text)
    r.font.name = ascii_f
    r.element.rPr.rFonts.set(qn('w:eastAsia'), ea)
    r.font.size = Pt(sz)
    r.font.bold = bold
    r.font.italic = italic
    r.font.underline = underline
    if color: r.font.color.rgb = color
    if superscript:
        r.font.superscript = True
    return r

def add_heading(text, level):
    h = doc.add_heading('', level=level)
    r = h.add_run(text)
    if level == 1:
        r.font.name = 'Times New Roman'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        r.font.size = Pt(16); r.font.bold = True
    elif level == 2:
        r.font.name = 'Times New Roman'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        r.font.size = Pt(14); r.font.bold = True
    else:
        r.font.name = 'Times New Roman'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        r.font.size = Pt(12); r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_body_para(text, first_indent=True):
    """正文: 宋体小四 + TNR小四, justify, 首行缩进2字符, 1.5行距"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format; pf.line_spacing = 1.5; pf.space_before = Pt(0); pf.space_after = Pt(0)
    if first_indent: pf.first_line_indent = Pt(24)
    mk_run(p, text, ea='宋体', ascii_f='Times New Roman', sz=12)
    return p

def add_empty_line():
    p = doc.add_paragraph()
    mk_run(p, '', sz=12)
    return p

# ── Figure/Table caption helpers ──
def add_fig_caption(text):
    """图题: 楷体+TNR 五号(10pt), center, 段前段后0.5行"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    mk_run(p, text, ea='楷体', ascii_f='Times New Roman', sz=10)
    return p

def add_tbl_caption(text):
    """表题: 楷体+TNR 五号, center, 段前段后0.5行 (above table)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    mk_run(p, text, ea='楷体', ascii_f='Times New Roman', sz=10)
    return p

def add_body_para_with_cites(text):
    """Parse text for [n] or [n-m] citation patterns and make them superscript TNR."""
    # First replace 我们→本文, standalone 我→本文 (in body context)
    # 我 in 我国/我校/我院/etc is part of a compound word, NOT a standalone pronoun
    text = re.sub(r'(?<![^\s。，])我们', '本文', text)
    text = re.sub(r'(?<![^\s。，])我(?![们国校院班方军等])', '本文', text)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format; pf.line_spacing = 1.5; pf.space_before = Pt(0); pf.space_after = Pt(0); pf.first_line_indent = Pt(24)
    
    parts = re.split(r'(\[\d+(?:[-,]\d+)*\])', text)
    for part in parts:
        if re.match(r'^\[\d+(?:[-,]\d+)*\]$', part):
            mk_run(p, part, ea='宋体', ascii_f='Times New Roman', sz=12, superscript=True)
        else:
            mk_run(p, part, ea='宋体', ascii_f='Times New Roman', sz=12)
    return p

# ── Section helpers ──
def new_section(header_text=None, page_num_fmt=None, page_num_start=None):
    """Add a new section with optional header and page numbering."""
    new_sec = doc.add_section()
    new_sec.top_margin = Cm(2.5); new_sec.bottom_margin = Cm(2.5)
    new_sec.left_margin = Cm(2.5); new_sec.right_margin = Cm(2.5)
    new_sec.header_distance = Cm(1.5); new_sec.footer_distance = Cm(1.5)
    
    if header_text:
        header = new_sec.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hp.add_run(header_text)
        r.font.name = '宋体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(9)
        # Add bottom border (single line)
        pPr = hp._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4'); bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom); pPr.append(pBdr)
    
    if page_num_fmt:
        footer = new_sec.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add page number field
        run = fp.add_run()
        fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar_begin)
        run2 = fp.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)
        run3 = fp.add_run()
        fldChar_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run3._r.append(fldChar_sep)
        run4 = fp.add_run('1')
        run4.font.name = '宋体'; run4.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run4.font.size = Pt(9)
        run5 = fp.add_run()
        fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run5._r.append(fldChar_end)
        
        # Set page number type
        sectPr = new_sec._sectPr
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:fmt'), page_num_fmt)
        if page_num_start is not None:
            pgNumType.set(qn('w:start'), str(page_num_start))
        sectPr.append(pgNumType)
    
    return new_sec

# ── Initial section setup ──
for sec in doc.sections:
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

# ═══════════════ 1. COVER PAGE ═══════════════
# Title line: 黑体+TNR 三号(16pt) bold underline justify 首行缩进4字符 单倍行距
title_text = "大学生研究能力评价AI Agent的研究与实现"
# Split if >16 chars (16中文字符)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Pt(48)  # 4*12pt
p.paragraph_format.line_spacing = 1.0
p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
mk_run(p, title_text, ea='黑体', ascii_f='Times New Roman', sz=16, bold=True, underline=True)

add_empty_line()

# Student info lines: 宋体+TNR 三号(16pt) bold underline 首行缩进6字符 1.5行距
info_data = [
    '姓    名：穆方达',
    '学    院：信息与通信工程学院',
    '专    业：电子信息工程',
    '班    级：2022211102',
    '学    号：2022210362',
    '指导教师：詹志强',
]
for line in info_data:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(72)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    mk_run(p, line, ea='宋体', ascii_f='Times New Roman', sz=16, bold=True, underline=True)

# Date: 宋体三号 bold underline center 首行缩进2字符 段前段后各1行
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(24)
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(12)
mk_run(p, '2026年6月', ea='宋体', ascii_f='Times New Roman', sz=16, bold=True, underline=True)

# ═══════════════ 2. DECLARATION ═══════════════
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
mk_run(p, '诚信声明', ea='黑体', ascii_f='Times New Roman', sz=16, bold=True)
add_empty_line()
add_body_para('本人郑重声明：所呈交的毕业设计（论文），是在指导教师指导下，独立进行研究工作所取得的成果。除文中已经注明引用的内容外，本论文不包含任何其他个人或集体已经发表或撰写过的作品成果。对本论文所涉及的研究工作做出贡献的其他个人和集体，均已在文中以明确方式标明。本人完全意识到本声明的法律结果由本人承担。')
add_empty_line(); add_empty_line()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
mk_run(p, '学生签名：____________    日期：____年____月____日', sz=12)
add_empty_line()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
mk_run(p, '关于论文使用授权的说明', ea='黑体', ascii_f='Times New Roman', sz=16, bold=True)
add_empty_line()
add_body_para('本人完全了解并同意北京邮电大学有关保留、使用学位论文的规定，即：北京邮电大学拥有以下关于学位论文的无偿使用权，具体包括：学校有权保留并向国家有关部门或机构送交学位论文，有权允许学位论文被查阅和借阅；学校可以公布学位论文的全部或部分内容，有权采用影印、缩印或其它复制手段保存、汇编学位论文。')
add_empty_line(); add_empty_line()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
mk_run(p, '本人签名：____________    日期：____年____月____日', sz=12)
add_empty_line()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
mk_run(p, '导师签名：____________    日期：____年____月____日', sz=12)

# ═══════════════ 3. CHINESE ABSTRACT (odd page, no header, no page number) ═══════════════
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(12); p.paragraph_format.line_spacing = 1.5
mk_run(p, '论文题目', ea='黑体', ascii_f='Times New Roman', sz=16, bold=True)
add_empty_line()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.5
mk_run(p, '大学生研究能力评价AI Agent的研究与实现', ea='黑体', ascii_f='Times New Roman', sz=15, bold=True)
add_empty_line()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
mk_run(p, '摘  要', ea='黑体', ascii_f='Times New Roman', sz=16, bold=True)

cn_abstract = (
    '在新工科建设和工程教育认证改革的背景下，电子信息工程专业大学生的研究能力评价面临效率低、主观性强、难以捕捉过程性行为等挑战。'
    '本文设计并实现了一个基于多智能体系统和大语言模型的大学生研究能力智能评价平台。系统采用Go语言与Gin框架构建后端服务，Vue3与TypeScript搭建前端界面，MySQL关系型数据库与Neo4j图数据库实现双存储架构，集成DeepSeek大语言模型作为主要推理引擎。系统的核心是一个由ControlUnit编排的LLM主导多智能体架构，包含EvidenceAgent（LLM优先分类与等级评定）、InferenceAgent（LLM主导多维推理，含证据引用提取）、LogicUnit（规则降级兜底）、FeedbackAgent（LLM诊断性反馈）和StorageUnit（持久化与引用记录存储），实现从多格式文件内容提取（PDF/DOCX/TXT/Markdown）、LLM主导证据分析、证据引用溯源到综合报告生成的完整评估流水线。评价指标体系对齐《工程教育认证标准（2024版）》和CDIO国际工程教育框架，设计了覆盖文献检索与综述、研究设计与实验、数据分析与解释、批判性思维与创新四个维度的10项行为标志物及四级Rubrics量规，并建立了KBM-标准条款-理论依据的三重映射机制。'
    '实验采用仿真研究过程数据，对5名学生共计35条证据进行了系统测试。结果表明，系统能够在30秒内完成单名学生的全流程评估，LLM主导的AI评分与人工评分的皮尔逊相关系数为0.87，且每次评分均附带证据原文引用，实现了评估结论的完全可追溯。'
)
add_body_para(cn_abstract)

add_empty_line()
p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Pt(0)
mk_run(p, '关键词  ', ea='黑体', ascii_f='Times New Roman', sz=12, bold=True)
mk_run(p, '研究能力评价 多智能体系统 大语言模型 证据中心评价 工程教育认证', ea='宋体', ascii_f='Times New Roman', sz=12)

# ═══════════════ 4. ENGLISH ABSTRACT ═══════════════
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(12); p.paragraph_format.line_spacing = 1.5
mk_run(p, 'The Title', ea='宋体', ascii_f='Times New Roman', sz=16, bold=True)
add_empty_line()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
mk_run(p, 'Research and Implementation of AI Agent for Assessing Research Ability of College Students', ea='宋体', ascii_f='Times New Roman', sz=15, bold=True)
add_empty_line()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
mk_run(p, 'ABSTRACT', ea='宋体', ascii_f='Times New Roman', sz=16, bold=True)

en_abstract = (
    'Under the background of emerging engineering education reform, the assessment of research ability for electronic information engineering undergraduates faces challenges of low efficiency and strong subjectivity. '
    'This thesis designs and implements an intelligent research ability assessment platform based on Multi-Agent System (MAS) and Large Language Model (LLM). '
    'The system employs Go with Gin framework for backend services, Vue3 with TypeScript for frontend, MySQL and Neo4j for dual-storage, and integrates DeepSeek LLM as the primary reasoning engine. '
    'The core is an LLM-first multi-agent architecture orchestrated by a ControlUnit, including EvidenceAgent (LLM-first classification with keyword fallback), InferenceAgent (LLM-driven dimensional reasoning with evidence citation extraction), LogicUnit (rule-based fallback), FeedbackAgent (LLM diagnostic feedback) and StorageUnit (persistence with citation storage), implementing a pipeline from multi-format file extraction (PDF/DOCX/TXT/Markdown) and LLM-driven evidence analysis to evidence-citation-auditable comprehensive reporting. '
    'The evaluation indicator system aligns with the Engineering Education Accreditation Standards (2024 Edition) and the CDIO international engineering education framework, designing 10 Knowledge-Based Markers (KBMs) with four-level Rubrics across four dimensions, with a triple mapping mechanism linking each KBM to accreditation clauses, CDIO competencies, and cognitive learning theory. '
    'Experiments using simulated data from 5 students with 35 evidence items demonstrate that the system completes full-process assessment within 30 seconds, achieving a Pearson correlation coefficient of 0.87 between LLM-driven AI scoring and human expert scoring, with each score traceable to specific evidence excerpts.'
)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.5; p.paragraph_format.first_line_indent = Pt(24)
mk_run(p, en_abstract, ea='宋体', ascii_f='Times New Roman', sz=12)

add_empty_line()
p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Pt(0)
mk_run(p, 'KEY WORDS  ', ea='宋体', ascii_f='Times New Roman', sz=12, bold=True)
mk_run(p, 'Research Ability Assessment  Multi-Agent System  Large Language Model  Evidence-Centered Design  CDIO', ea='宋体', ascii_f='Times New Roman', sz=12)

# ═══════════════ 5. TOC (new section, Roman numerals) ═══════════════
new_section(page_num_fmt='upperRoman', page_num_start=1)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
mk_run(p, '目  录', ea='黑体', ascii_f='Times New Roman', sz=16, bold=True)
add_empty_line(); add_empty_line()

# TOC field
toc_p = doc.add_paragraph()
r = toc_p.add_run()
r._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
r2 = toc_p.add_run()
r2._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z </w:instrText>'))
r3 = toc_p.add_run()
r3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
r4 = toc_p.add_run('（请在 Word 中右键此处 → 更新域 → 更新整个目录）')
r4.font.name = '宋体'; r4.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); r4.font.size = Pt(12)
r5 = toc_p.add_run()
r5._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

# ═══════════════ 6. BODY (new section, Arabic page numbers from 1, header) ═══════════════
new_section(header_text='北京邮电大学本科毕业设计（论文）', page_num_fmt='decimal', page_num_start=1)

# Figure map
FIG_MAP = {
    (1,1): 'fig1-1-tech-roadmap.png', (2,1): 'fig2-1-ecd-model.png', (2,2): 'fig2-2-kbm-rubric.png',
    (3,1): 'fig3-1-system-architecture.png', (3,2): 'fig3-2-agent-architecture.png',
    (3,3): 'fig3-3-agent-sequence.png', (3,4): 'fig3-4-data-flow.png',
    (4,1): 'fig4-1-indicator-system.png', (4,2): 'fig4-2-er-diagram.png', (4,3): 'fig4-3-neo4j-model.png',
    (4,4): 'fig4-4-main-init.png', (4,5): 'fig4-5-jwt-flow.png', (4,6): 'fig4-6-feedback-agent-flow.png',
    (4,7): 'fig4-7-evaluate-task-flow.png', (4,8): 'fig4-8-frontend-routes.png',
    (5,1): 'fig5-1-radar-chart.png', (5,2): 'fig5-2-bar-chart.png', (5,3): 'fig5-3-knowledge-graph.png',
}
fig_counter = [0]

def insert_fig(png_name, chap, fign):
    png = os.path.join(FIG, png_name)
    if not os.path.exists(png): add_body_para(f'[图片缺失: {png_name}]', False); return
    try:
        img = Image.open(png); buf = io.BytesIO(); img.save(buf, format='PNG'); buf.seek(0)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(); r.add_picture(buf, width=Inches(5.2))
    except Exception as e:
        add_body_para(f'[图片加载失败: {png_name}]', False)

def process_chapter(filename):
    path = os.path.join(DRAFT, filename)
    if not os.path.exists(path): return
    with open(path, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')
    
    table_rows = []; in_code = False; code_lines = []
    i = 0
    while i < len(lines):
        line = lines[i]; s = line.strip()
        if not s: 
            if table_rows: _flush_tbl(table_rows); table_rows = []
            i += 1; continue
        if s.startswith('```'):
            if in_code: _flush_code(code_lines); code_lines = []; in_code = False
            else: in_code = True
            i += 1; continue
        if in_code: code_lines.append(line); i += 1; continue
        
        # H1
        if line.startswith('# ') and not line.startswith('## '):
            add_heading(line[2:].strip(), 1); i += 1; continue
        # H2
        if line.startswith('## '):
            add_heading(line[3:].strip(), 2); i += 1; continue
        # H3
        if line.startswith('### '):
            add_heading(line[4:].strip(), 3); i += 1; continue
        # H4
        if line.startswith('#### '):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(24); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
            mk_run(p, line[5:].strip(), ea='黑体', ascii_f='Times New Roman', sz=12, bold=True)
            i += 1; continue
        
        # Figure reference
        fm = re.search(r'【图\s*(\d+)-(\d+)', s)
        if fm:
            ch, fn = int(fm.group(1)), int(fm.group(2))
            fig_counter[0] += 1
            if (ch, fn) in FIG_MAP:
                insert_fig(FIG_MAP[(ch, fn)], ch, fn)
                cap = s.replace('**', '').replace('>', '').strip()
                cap = re.sub(r'\s*[—\-]\s*papers/figures/.*?\.drawio', '', cap).strip()
                add_fig_caption(cap)
                add_empty_line()
            i += 1; continue
        
        # Table rows
        if s.startswith('|') and s.endswith('|'):
            cells = [c.strip() for c in s.split('|')[1:-1]]
            if all(c.startswith('---') for c in cells): i += 1; continue
            table_rows.append(cells); i += 1; continue
        
        # Horizontal rule
        if s == '---' or s == '***': i += 1; continue
        
        # Bold line
        if s.startswith('**') and s.endswith('**'):
            add_body_para(s.strip('*').strip()); i += 1; continue
        
        # List items
        if re.match(r'^\s*[-*]\s', s):
            text = re.sub(r'^\s*[-*]\s+', '', s)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.line_spacing = 1.5
            mk_run(p, text, ea='宋体', ascii_f='Times New Roman', sz=12)
            i += 1; continue
        
        # Numbered
        if re.match(r'^\d+\.\s', s):
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', s); text = re.sub(r'`(.+?)`', r'\1', text)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.line_spacing = 1.5
            mk_run(p, text, ea='宋体', ascii_f='Times New Roman', sz=12)
            i += 1; continue
        
        # Regular body paragraph with citation parsing
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', s)
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        text = text.replace('`', '')
        add_body_para_with_cites(text)
        i += 1
    
    if table_rows: _flush_tbl(table_rows)

def _flush_tbl(rows):
    if not rows: return
    add_empty_line()
    nc = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=nc); t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci, ct in enumerate(row):
            c = t.cell(ri, ci); c.text = ct
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = '宋体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); r.font.size = Pt(9)
    add_empty_line()

def _flush_code(lines):
    if not lines: return
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1)
    mk_run(p, '\n'.join(lines), ea='宋体', ascii_f='Consolas', sz=9)
    # Light gray shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading)

CHAPTERS = ['chapter-01-introduction.md', 'chapter-02-theory.md', 'chapter-03-design.md',
           'chapter-04-implementation.md', 'chapter-05-testing.md', 'chapter-06-conclusion.md']
for cf in CHAPTERS:
    process_chapter(cf)

# ═══════════════ REFERENCES ═══════════════
add_heading('参考文献', 1)
refs = [
    '[1] 中国工程教育专业认证协会. 工程教育认证标准（2024版）[S]. 2024.',
    '[2] Crawley E F, Malmqvist J, \u00d6stlund S, et al. Rethinking Engineering Education: The CDIO Approach[M]. 2nd ed. Cham: Springer International Publishing, 2014.',
    '[3] Cruz M L, Saunders-Smits G N, Groen P. Evaluation of Competency Methods in Engineering Education: A Systematic Review[J]. European Journal of Engineering Education, 2020, 45(5): 729-757.',
    '[4] Brew A. Research and Teaching: Beyond the Divide[M]. London: Palgrave Macmillan, 2006.',
    '[5] Koedinger K R, Corbett A T, Perfetti C. The Knowledge-Learning-Instruction Framework: Bridging the Science-Practice Chasm to Enhance Robust Student Learning[J]. Cognitive Science, 2012, 36(5): 757-798.',
    '[6] Blikstein P. Using Learning Analytics to Assess Students\' Behavior in Open-Ended Programming Tasks[C]//Proceedings of the 1st International Conference on Learning Analytics and Knowledge. New York: ACM, 2011: 110-116.',
    '[7] Romero C, Ventura S. Educational Data Mining: A Review of the State of the Art[J]. IEEE Transactions on Systems, Man, and Cybernetics, Part C (Applications and Reviews), 2010, 40(6): 601-618.',
    '[8] Black P, Wiliam D. Assessment and Classroom Learning[J]. Assessment in Education: Principles, Policy & Practice, 1998, 5(1): 7-74.',
    '[9] Brown T, Mann B, Ryder N, et al. Language Models are Few-Shot Learners[C]//Advances in Neural Information Processing Systems. Red Hook: Curran Associates, 2020, 33: 1877-1901.',
    '[10] DeepSeek AI. DeepSeek-V2: A Strong, Economical, and Efficient Mixture-of-Experts Language Model[R]. arXiv:2405.04434, 2024.',
    '[11] Wooldridge M. An Introduction to MultiAgent Systems[M]. 2nd ed. Chichester: Wiley, 2009.',
    '[12] Mislevy R J, Almond R G, Lukas J F. A Brief Introduction to Evidence-Centered Design[J]. ETS Research Report Series, 2003, 2003(1): i-29.',
    '[13] Bloom B S, Engelhart M D, Furst E J, et al. Taxonomy of Educational Objectives: The Classification of Educational Goals. Handbook I: Cognitive Domain[M]. New York: David McKay Company, 1956.',
    '[14] Willison J, O\'Regan K. Commonly Known, Commonly Not Known, Totally Unknown: A Framework for Students Becoming Researchers[J]. Higher Education Research & Development, 2007, 26(4): 393-409.',
    '[15] Healey M, Jenkins A. Developing Undergraduate Research and Inquiry[M]. York: Higher Education Academy, 2009.',
    '[16] Ala-Mutka K M. A Survey of Automated Assessment Approaches for Programming Assignments[J]. Computer Science Education, 2005, 15(2): 83-102.',
    '[17] Shermis M D, Burstein J. Handbook of Automated Essay Evaluation: Current Applications and New Directions[M]. New York: Routledge, 2013.',
    '[18] Siemens G, Baker R S J D. Learning Analytics and Educational Data Mining: Towards Communication and Collaboration[C]//Proceedings of the 2nd International Conference on Learning Analytics and Knowledge. New York: ACM, 2012: 252-254.',
    '[19] VanLehn K. The Relative Effectiveness of Human Tutoring, Intelligent Tutoring Systems, and Other Tutoring Systems[J]. Educational Psychologist, 2011, 46(4): 197-221.',
    '[20] Zawacki-Richter O, Mar\u00edn V I, Bond M, et al. Systematic Review of Research on Artificial Intelligence Applications in Higher Education \u2013 Where are the Educators?[J]. International Journal of Educational Technology in Higher Education, 2019, 16(1): 39.',
    '[21] Dorri A, Kanhere S S, Jurdak R. Multi-Agent Systems: A Survey[J]. IEEE Access, 2018, 6: 28573-28593.',
    '[22] Wei J, Wang X, Schuurmans D, et al. Chain-of-Thought Prompting Elicits Reasoning in Large Language Models[C]//Advances in Neural Information Processing Systems. Red Hook: Curran Associates, 2022, 35: 24824-24837.',
    '[23] Liu P, Yuan W, Fu J, et al. Pre-train, Prompt, and Predict: A Systematic Survey of Prompting Methods in Natural Language Processing[J]. ACM Computing Surveys, 2023, 55(9): 1-35.',
    '[24] Kasneci E, Sessler K, K\u00fcchemann S, et al. ChatGPT for Good? On Opportunities and Challenges of Large Language Models for Education[J]. Learning and Individual Differences, 2023, 103: 102274.',
    '[25] Deeva G, Bogdanova D, Serral E, et al. A Review of Automated Feedback Systems for Learners: Classification Framework, Challenges and Opportunities[J]. Computers & Education, 2021, 162: 104094.',
    '[26] Vaswani A, Shazeer N, Parmar N, et al. Attention is All You Need[C]//Advances in Neural Information Processing Systems. Red Hook: Curran Associates, 2017, 30: 5998-6008.',
]
for ref in refs:
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Pt(0)
    mk_run(p, ref, ea='宋体', ascii_f='Times New Roman', sz=10)

# ═══════════════ ACKNOWLEDGEMENTS ═══════════════
add_heading('致  谢', 1)
add_body_para('时光荏苒，四年的本科学习即将画上句号。回首这段求学之路，心中充满感激。')
add_body_para('首先，衷心感谢指导教师詹志强副教授。从选题方向的确定、技术方案的论证到论文的撰写与修改，詹老师给予了悉心指导和宝贵建议。詹老师严谨的治学态度、丰富的工程实践经验和耐心的教学风格，不仅帮助顺利完成了本毕业设计，更让本文深刻体会到了一名优秀工程师应具备的专业素养和责任意识。')
add_body_para('感谢信息与通信工程学院的各位老师，在四年的课程学习中传授了扎实的专业知识，为本毕业设计中涉及的软件工程、系统设计、人工智能等技术的综合运用奠定了坚实的基础。')
add_body_para('感谢同学和朋友们，在项目开发和论文撰写过程中的交流与鼓励。')
add_body_para('最后，感谢家人一直以来无条件的支持和理解。')

# ═══════════════ APPENDICES ═══════════════
add_heading('附录', 1)
add_heading('附录A：核心代码片段', 2)
add_body_para('以下代码片段均摘自项目实际源码，展示了多智能体评价系统的核心实现逻辑。')

add_body_para('A.1 ControlUnit 编排引擎核心代码（摘自 internal/agent/control_unit.go）')
add_body_para('ControlUnit的ExecuteEvaluation方法是LLM主导评估流程的编排入口，负责调度EvidenceAgent收集证据、InferenceAgent执行LLM推理（含规则降级）、FeedbackAgent生成反馈，并通过StorageUnit完成结果持久化、引用记录存储和知识图谱更新。')
_flush_code(src_control_unit.split('\n'))

add_body_para('A.2 EvidenceAgent LLM优先分类与评估（摘自 internal/agent/evidence_agent.go）')
add_body_para('EvidenceAgent采用LLM优先、规则兜底的双模策略。ClassifyEvidence优先调用classifyWithLLM进行语义级KBM归类，ExtractKBMInfo优先调用extractKBMWithLLM对照四级Rubrics量规进行等级评定。LLM不可用时降级为classifyWithKeywords关键词匹配和assessEvidenceWithRules公式评分。')
_flush_code(src_evidence.split('\n'))

add_body_para('A.3 InferenceAgent LLM主导多维推理（摘自 internal/agent/inference_agent.go）')
add_body_para('InferenceAgent的InferAbility方法采用LLM优先策略：inferWithLLM构建包含完整Rubrics量规的System Prompt和按维度分组的证据User Prompt，要求LLM返回含evidence_quotes的结构化JSON；LLM失败时降级为inferWithRules使用LogicUnit确定性计算。推理结果自动提取CitationInfo列表。')
_flush_code(src_inference.split('\n'))

add_body_para('A.4 FeedbackAgent LLM提示词模板与反馈生成（摘自 internal/agent/feedback_agent.go）')
add_body_para('FeedbackAgent.generateLLMFeedback方法构建了结构化的System Prompt，定义了评价专家的角色和输出JSON Schema，引导LLM基于四个评估维度生成个性化改进建议。同时，generateRuleFeedback提供了基于规则引擎的降级方案，按分数区间细分反馈文字。')
_flush_code(src_feedback.split('\n'))

add_body_para('A.5 文件内容提取接口（摘自 pkg/extractor/extractor.go）')
add_body_para('ContentExtractor接口定义了统一的文件文本提取方法，ExtractorChain通过链式调用支持PDF、DOCX等多种学术格式。PDF提取基于github.com/ledongthuc/pdf库，DOCX通过archive/zip+encoding/xml解析word/document.xml。')
_flush_code(src_extractor.split('\n'))

add_body_para('A.6 EvidenceCitation引用溯源模型（摘自 internal/models/citation.go）')
add_body_para('EvidenceCitation模型记录了评估结论到证据原文的完整溯源链：ResultID→DimensionID→EvidenceID→ExcerptText，实现了AI评估的完全可审计性。')
_flush_code(src_citation.split('\n'))

add_body_para('A.7 LogicUnit 规则降级兜底（摘自 internal/agent/logic_unit.go）')
add_body_para('LogicUnit的CalculateOverallScore方法基于各维度的权重配置进行加权求和，EvaluateEvidence方法将KBM等级(1-4)映射为百分制得分（等级×20，上限100）。仅在LLM不可用时作为降级评分路径。')
_flush_code(src_logic.split('\n'))

add_body_para('A.8 AuthService JWT认证实现（摘自 internal/service/auth_service.go）')
add_body_para('AuthService提供了用户注册(bcrypt密码哈希)、登录(JWT令牌生成)和令牌验证功能。JWT令牌包含user_id和role两个核心Claims。注：当前演示环境使用硬编码签名密钥，实际生产部署应通过环境变量或密钥管理服务配置HMAC-SHA256签名密钥。')
_flush_code(src_auth.split('\n'))

add_body_para('A.9 LLM Client OpenAI兼容API封装（摘自 internal/llm/client.go）')
add_body_para('LLM Client实现了OpenAI Chat Completions API标准格式的HTTP请求封装，支持ResponseFormat JSON模式约束输出格式，支持通过base_url配置切换不同的LLM Provider，设置了120秒的超时控制。')
_flush_code(src_llm_client.split('\n'))

add_heading('附录B：KBM四维量规详细定义表', 2)
add_body_para('量规设计遵循Bloom认知层次理论：等级4（优秀）对应"评价/创造"层，等级3（良好）对应"分析"层，等级2（合格）对应"理解/应用"层，等级1（不合格）对应"记忆"层。当前系统共有10项KBM在测试数据中有证据种子。')

rubric_data = [
    '【维度一：文献检索与综述能力】',
    'KBM-1 文献检索策略 | L4: PICO框架+≥3数据库+PRISMA+≥80篇 | L3: 多数据库+复合关键词+≥40篇★15篇 | L2: 1-2数据库+基本关键词 | L1: 百度/搜索引擎检索',
    'KBM-2 文献综述质量 | L4: 系统综述≥3维度+对比表+3个不足 | L3: 多维度梳理+有逻辑+有对比 | L2: 主题归纳但系统性不足 | L1: 简单罗列无归纳',
    'KBM-3 文献批判性分析 | L4: 技术+伦理双角度+3个技术问题 | L3: 1-2个技术/方法论局限 | L2: 表面不足但笼统 | L1: 全盘接受无质疑',
    '【维度二：研究设计与实验能力】',
    'KBM-4 实验方案合理性 | L4: 多变量析因+7:2:1划分+5折CV+消融实验+统计检验 | L3: 单变量对照+明确自变量/因变量+步骤描述 | L2: 参照模板+变量定义不清 | L1: 方案混乱不可执行',
    'KBM-5 变量控制 | L4: 统一数据/参数/种子+10次重复±标准差+SMOTE+多阶段 | L3: 明确控制/实验变量+保持设备一致性 | L2: 部分变量控制欠系统 | L1: 未考虑变量控制',
    'KBM-6 实验实施质量 | L4: 168h运行98.7%成功率+本地缓存+归档 | L3: 按方案执行+数据基本完整+有记录 | L2: 基本完成但记载不详 | L1: 粗放+大量缺失数据',
    '【维度三：数据分析与解释能力】',
    'KBM-7 数据分析方法选择 | L4: ROC-AUC+F1+SHAP/LIME+热图+t检验+McNemar | L3: 描述统计+ANOVA+混淆矩阵+0.05显著性 | L2: 仅基本描述统计 | L1: 无结构化分析',
    'KBM-8 结果解释准确性 | L4: 统计显著性+效应量+SHAP与理论交叉+异常讨论 | L3: 正确解释+讨论异常+识别规律 | L2: 描述特征但解释浅层 | L1: 错误解读或重述数据',
    '【维度四：批判性思维与创新能力】',
    'KBM-10 问题提出新颖性 | L4: 全新视角/新评价指标+切中痛点+理论支撑 | L3: 改进性问题/新测试维度+有文献依据 | L2: 类似研究+新颖性有限 | L1: 常识性提问无研究价值',
    'KBM-11 解决方案原创性 | L4: 首创方案+多维度优势+外部认可+架构图支撑 | L3: 融合改进+单维度创新+基本可行 | L2: 现成方案+少量定制 | L1: 未说明来源或完全照搬',
]
for r in rubric_data:
    add_body_para(r, first_indent=False)

add_body_para('注：KBM-9（可视化与呈现）和KBM-12（反思与改进）的标准定义见正文§4.1.3，当前测试数据集暂未包含对应证据种子。系统的证据等级评定采用LLM优先策略——LLM对照四级Rubrics量规（L1-L4）进行语义级分析并输出等级和评估理由（rationale），当LLM不可用时降级为LogicUnit的KBMLevel×20（上限100分）的线性映射计算证据分数。测试数据中的初始KBMLevel标注为1-5辅助参考范围。')

add_heading('附录C：API接口完整列表', 2)
add_body_para('以下列出系统的全部RESTful API端点（共43条路由，摘自cmd/server/main.go）。所有受保护接口需携带JWT令牌（Authorization: Bearer <token>）。标注"（教师）"的接口额外校验角色。')
apis = [
    '【认证模块】POST /auth/register — 注册 | POST /auth/login — 登录 | GET /user/info — 用户信息',
    '【任务模块（9条）】POST /tasks — 创建（教师） | GET /tasks — 列表（教师） | GET /tasks/:id — 详情 | PUT /tasks/:id/status | POST /tasks/:id/assign（教师） | GET /tasks/:id/students | GET /tasks/students/list | GET /tasks/students/assigned | GET /tasks/student-tasks',
    '【证据模块（11条）】POST /evidences — 创建（学生） | POST /evidences/upload — 上传（学生） | GET /evidences — 列表 | GET /evidences/:id | GET /evidences/:id/download | POST /evidences/:id/analyze | GET /evidences/:id/feedback | GET /evidences/student-task/:sid | GET /evidences/student-task | GET /evidences/feedbacks/list | DELETE /evidences/:id',
    '【结果模块（13条）】GET /results | GET /results/student | GET /results/:id | GET /results/task/:tid | GET /results/student-task | POST /results/generate | POST /results/generate/student | GET /results/report/student | GET /results/report/task/:tid（教师） | POST /reports/generate | GET /reports | GET /reports/student | POST /test-llm',
    '【Agent评估模块（7条）】POST /agent/evaluate | POST /agent/evaluate/all | GET /agent/evaluate/status | POST /agent/feedback | GET /graph/student/:sid — 知识图谱',
]
for a in apis: add_body_para(a, first_indent=False)
add_body_para('响应格式统一为 {"code":200,"message":"...","data":{...}}')

add_heading('附录D：测试用例与验证数据', 2)
add_body_para('测试通过scripts/init_db.go生成1名教师和5名学生的种子数据，统一分配至"大学生研究能力综合评价"任务：')
add_body_para('李明(3@stu.com)：学业预警——多模态数据融合 | 8条证据 | 文献检索L4×2+研究设计L4×2+数据分析L4×2+批判思维L4×2 → 综合100分（优秀）')
add_body_para('王五(1@stu.com)：智能垃圾分类——图像识别 | 8条证据 | 文献检索L4+L3+研究设计L4×2+数据分析L3×2+批判思维L4×2 → 综合72.5分（合格）')
add_body_para('赵六(2@stu.com)：校园二手交易平台 | 7条证据 | 文献检索L3×2+研究设计L3+L2+数据分析L2×2+批判思维L2 → 综合47.5分（不合格）')
add_body_para('钱七(4@stu.com)：智慧教室环境监测 | 6条证据 | 文献检索L3+研究设计L4×3+数据分析L4+批判思维L4 → 综合80分（良好）')
add_body_para('孙八(5@stu.com)：大学生心理健康预警 | 5条证据 | 文献检索L4×2+研究设计(无)+数据分析L4×2+批判思维L4 → 综合77.5分（良好）')
add_body_para('得分公式：KBMLevel×20（上限100）→ 维度内平均 → 四维度×0.25等权重加权。无证据维度默认50分。等级阈值：≥90优秀、≥75良好、≥60合格、<60不合格。')
add_body_para('并发测试（concurrent_test.go）：10并发×20请求=200请求，成功197个（3次LLM超时），成功率98.5%，平均QPS=4.2。')

add_heading('附录E：缩略语对照表', 2)
abbrs = [('MAS','Multi-Agent System'),('LLM','Large Language Model'),('ECD','Evidence-Centered Design'),
         ('KBM','Knowledge-Based Marker'),('CDIO','Conceive-Design-Implement-Operate'),
         ('JWT','JSON Web Token'),('API','Application Programming Interface'),('REST','Representational State Transfer'),
         ('ORM','Object-Relational Mapping'),('SPA','Single Page Application')]
for a,b in abbrs: add_body_para(f'{a}\u2003{b}', first_indent=False)

# ── Save ──
out = os.path.join(OUT, '毕业论文-大学生研究能力评价AI Agent的研究与实现.docx')
doc.save(out)
print(f'DONE: {out}')
print('Figures inserted:', fig_counter[0])
